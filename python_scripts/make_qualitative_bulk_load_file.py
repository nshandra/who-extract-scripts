import json
import os
import sys
import argparse
from docx import Document
from collections import namedtuple
import openpyxl
from openpyxl.cell import MergedCell


Metadata_ids = namedtuple("Metadata_ids", "sections, data_elements, countries, combos")


def get_country_id(country, countries_ids):
    for key in countries_ids:
        if country in key:
            country_key = key

    if country_key:
        return countries_ids[country_key]
    else:
        raise ValueError(f'Can\'t find id for: {country}')


def get_data_element_id(de, data_elements_ids):
    if de in data_elements_ids:
        return data_elements_ids[de]
    else:
        raise ValueError(f'Can\'t find id for: {de}')

def get_metadata_ids(workbook):
    sections_id_dict = {}
    data_elements_id_dict = {}
    countries_id_dict = {}
    combos_id_dict = {}
    sheet = workbook['Metadata']

    for row in sheet.iter_rows(min_row=2, min_col=1, max_col=5, values_only=True):
        identifier = row[0]
        type_col = row[1]
        name = str(row[2]).strip()
        Option_set = row[4] if row[4] else False

        
        if type_col == 'sections':
            sections_id_dict[name] = identifier

        if type_col == 'categoryOptionCombos':
            combos_id_dict[name] = identifier

        if type_col == 'dataElements' and not Option_set:
            data_elements_id_dict[name] = identifier

        if type_col == 'organisationUnit':
            countries_id_dict[name] = identifier

    return Metadata_ids(sections_id_dict, data_elements_id_dict, countries_id_dict, combos_id_dict)


def make_matched_values(all_tables_data, ids: Metadata_ids):
    country_id = get_country_id(COUNTRY, ids.countries)
    default_combo_id = ids.combos['default']

    data = {}
    data[country_id] = {}
    data[country_id][YEAR] = {}

    for table_data in all_tables_data:
        for de_name, value in table_data.items():
            if not de_name in ids.sections:
                de_id = get_data_element_id(de_name, ids.data_elements)
                if de_id not in data[country_id][YEAR]:
                    data[country_id][YEAR][de_id] = {}

            data[country_id][YEAR][de_id][default_combo_id] = value

    return data


def write_org_unit(last_cell, matched_values):
    for country_id, country_data in matched_values.items():
        for year in country_data:
            new_cell = last_cell.offset(row=1, column=0)
            new_cell.value = f'=_{country_id}'

            last_cell = new_cell

    return last_cell


def write_years(last_cell, matched_values):
    for country_id, country_data in matched_values.items():
        for year in country_data:
            new_cell = last_cell.offset(row=1, column=0)
            new_cell.value = year

            last_cell = new_cell

    return last_cell


def write_data(col_indicator, col_combo, last_cell, matched_values):

    for country_id, country_data in matched_values.items():
        for year, data_elements in country_data.items():
            for indicator_id, indicator_combos in data_elements.items():
                if indicator_id == col_indicator:
                    for combo_id, value in indicator_combos.items():
                        ids = combo_id.split(
                            '|') if '|' in combo_id else combo_id
                        if col_combo in ids or (col_combo == 'HllvX50cXC0' and combo_id == 'gEWtgad4feW'):
                            new_cell = last_cell.offset(row=1, column=0)
                            new_cell.value = value

                            last_cell = new_cell

    return last_cell


def write_values(workbook, matched_values):
    sheet = workbook['Data Entry']
    workbook.active = workbook['Data Entry']

    for index, col in enumerate(sheet.iter_cols(min_row=4)):
        if index == 0:
            last_cell = col[-1]
            write_org_unit(last_cell, matched_values)
        if index == 1:
            last_cell = col[-1]
            write_years(last_cell, matched_values)
        if index == 2:
            pass
        if index > 2:
            if not isinstance(col[0], MergedCell):
                col_indicator = str(col[0].value).split('=_')[-1]
            col_combo = str(col[1].value).split('=_')[-1]
            last_cell = col[-1]

            write_data(col_indicator, col_combo,
                            last_cell, matched_values)

    workbook.save(OUT_FILENAME)


def get_country_and_year(document):
    """
    Reads the first line from the source DOCX file and returns the metadata country and year.

    :param document: The document loaded from the source DOCX file.
    :type document: Document
    :return: The country and year of the data in the DOCX file.
    :rtype: tuple (str, str)
    """

    first_paragraph = document.paragraphs[0]
    first_line = first_paragraph.text

    country, year = first_line.replace(")", "").split(" (")

    if not year.isnumeric() or len(year) != 4:
        raise ValueError(f'Invalid year format for: {year}')

    return country, year


def extract_longtext_tables(document):
    """
    Extracts tables from the source DOCX file and returns them as a list of dictionaries.

    :param document: The document loaded from the source DOCX file.
    :type document: Document
    :return: A list of dictionaries representing the tables in the .docx file.
    :rtype: list
    """

    tables_data_list = []

    for table in document.tables:
        if len(table.columns) != 2:
            continue

        table_data = {}
        for row in table.rows:
            key = row.cells[0].text.rstrip()
            value = row.cells[1].text.rstrip()
            if key and value:
                table_data[key] = value
            else:
                debug(f'Empty row with {"DE: " + key if key else ""}{" and " if key and value else ""}{"value: " + value if value else ""} in document')

        tables_data_list.append(table_data)

    return tables_data_list


def debug(*msg):
    if DEBUG:
        print(*msg)


def filepath_exists(filepath):
    return True if os.path.isfile(filepath) else False


def main():
    """
    Parses command-line arguments and extracts tables from a .docx file.
    """
    parser = argparse.ArgumentParser(description='Process DOCX files into "Bulk Load" XLSX files. \
                                     The script needs a template, it either can be supplied with the --xlsx_template \
                                     argument or by placing a template named "Qualitative_Data_UHCPW_Template.xlsx" \
                                     in the same folder as the script.\
                                     Outputs to a XLSX file named <COUNTRY>_<YEAR>_Qualitative_Data.xlsx.')
    parser.add_argument('docx_filename', type=str,
                        help='The path to the DOCX source file.')
    parser.add_argument('-x', '--xlsx_template', type=str,
                        help='Bulk Load Qualitative XLSX template file path, \
                            if empty tries to open "Qualitative_Data_UHCPW_Template.xlsx"')
    parser.add_argument('-d', '--debug', action='store_true',
                        help='Display debug logs, its recommended to redirect the output into a file, e.g: ... > log.txt')
    args = parser.parse_args()

    if not filepath_exists(args.docx_filename):
        parser.error(f'The source file: {args.docx_filename} doesn\'t exist')

    global OUT_FILENAME, DEFAULT_TEMPLATE, DEBUG, COUNTRY, YEAR
    DEFAULT_TEMPLATE = 'Qualitative_Data_UHCPW_Template.xlsx'
    DEBUG = args.debug

    if not args.xlsx_template and filepath_exists(DEFAULT_TEMPLATE):
        args.xlsx_template = DEFAULT_TEMPLATE
    elif not filepath_exists(args.xlsx_template):
        parser.error(f'The template: {args.xlsx_template} doesn\'t exist')
    else:
        parser.error(
            f'The default template: {DEFAULT_TEMPLATE} doesn\'t exist')

    document = Document(args.docx_filename)

    COUNTRY, YEAR = get_country_and_year(document)
    OUT_FILENAME = f'{COUNTRY}_{YEAR}_Qualitative_Data.xlsx'
    debug(COUNTRY, YEAR)

    longtext_tables_data = extract_longtext_tables(document)
    debug(json.dumps(longtext_tables_data))

    try:
        wb = openpyxl.load_workbook(args.xlsx_template)
    except Exception as e:
        print(e)
        sys.exit(1)

    ids = get_metadata_ids(wb)

    debug(f'sections ids:\n len: {len(ids.sections)}\n values:\n', json.dumps(ids.sections))
    debug(f'data_elements ids:\n len: {len(ids.data_elements)}\n values:\n', json.dumps(ids.data_elements))
    debug(f'countries ids:\n len: {len(ids.countries)}\n values:\n', json.dumps(ids.countries))
    debug(f'combos ids:\n len: {len(ids.combos)}\n values:\n', json.dumps(ids.combos))

    matched_values = make_matched_values(longtext_tables_data, ids)

    debug(f'matched_values:\n', json.dumps(matched_values))

    write_values(wb, matched_values)


if __name__ == '__main__':
    main()
